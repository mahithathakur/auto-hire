import os
import pdfplumber
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# ---------- JD Extraction ----------

def get_jd_text():
    jd_docx_path = os.path.join("downloads", "jd.docx")
    if os.path.exists(jd_docx_path):
        try:
            doc = docx.Document(jd_docx_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            return full_text
        except Exception as e:
            print(f"❌ Error reading JD file: {e}")
            return ""
    else:
        print("⚠️ JD file not found at downloads/jd.docx")
        return ""

# ---------- Resume Text Extraction ----------

def extract_text_from_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ''
            for page in pdf.pages:
                text += page.extract_text() or ''
        return text
    except Exception as e:
        print(f"❌ Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        return ""

# ---------- JD Matching ----------

def calculate_score(resume_text, jd_text):
    try:
        documents = [resume_text, jd_text]
        vectorizer = TfidfVectorizer(stop_words="english")
        tfidf_matrix = vectorizer.fit_transform(documents)
        score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        return score
    except Exception as e:
        print(f"❌ Error calculating JD score: {e}")
        return 0.0
